import streamlit as st
import re
import json
from fpdf import FPDF
from docx import Document
from datetime import datetime
import io

# --- Helper Class for PDF Generation ---
class ReportPDF(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 12)
        self.cell(0, 10, 'AI Compliance Co-Pilot Report', 0, 1, 'C')
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 5, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1, 'C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

# --- File Generation Functions ---
def _generate_report_pdf(report_markdown: str, title: str) -> bytes:
    """Generates a PDF document from a markdown report."""
    pdf = ReportPDF()
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 10, title, 0, 1, 'L')
    pdf.ln(5)

    findings = re.split(r'(\[Risk - High\]|\[Warning - Medium\]|\[Pass\])', report_markdown)
    
    pdf.set_font('Helvetica', '', 11)
    for i in range(1, len(findings), 2):
        tag = findings[i]
        content = findings[i+1].strip()
        
        pdf.set_font('Helvetica', 'B', 12)
        if "Risk" in tag: pdf.set_text_color(220, 53, 69) # Red
        elif "Warning" in tag: pdf.set_text_color(255, 193, 7) # Yellow/Amber
        elif "Pass" in tag: pdf.set_text_color(25, 135, 84) # Green
        pdf.multi_cell(0, 7, f"{tag}\n{content.splitlines()[0]}")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 7, "\n".join(content.splitlines()[1:]).strip())
        pdf.ln(5)

    return pdf.output(dest='S').encode('latin-1')

def _generate_report_docx(report_markdown: str, title: str) -> bytes:
    """Generates a DOCX document from a markdown report."""
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d")}')
    # Simple conversion, can be made more sophisticated
    cleaned_text = report_markdown.replace('**', '')
    document.add_paragraph(cleaned_text)
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _generate_report_json(report_markdown: str) -> bytes:
    """Converts a markdown report to a structured JSON object."""
    findings_list = []
    findings = re.split(r'(\[Risk - High\]|\[Warning - Medium\]|\[Pass\])', report_markdown)
    
    for i in range(1, len(findings), 2):
        tag = findings[i].strip("[]")
        content = findings[i+1].strip()
        findings_list.append({"status": tag, "details": content})
        
    return json.dumps(findings_list, indent=4).encode('utf-8')

# --- Main UI Function ---
def handle_report_display_and_download(report_markdown: str, base_filename: str):
    """
    Renders the styled compliance report and provides a full suite of download buttons.
    """
    # Professional header
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #1e40af 0%, #1e3a8a 100%);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        color: white;
        text-align: center;
        box-shadow: 0 10px 25px rgba(0, 0, 0, 0.1);
    ">
        <h2 style="margin: 0; font-size: 1.75rem; font-weight: 700;">Compliance Analysis Report</h2>
        <p style="margin: 0.5rem 0 0 0; font-size: 1rem; opacity: 0.9;">
            Comprehensive regulatory compliance assessment with detailed findings and recommendations
        </p>
    </div>
    """, unsafe_allow_html=True)

    if not report_markdown:
        st.info("No compliance findings to display.")
        return
    
    # Enhanced report display with proper formatting
    st.markdown("""
    <style>
    .report-finding { 
        border-radius: 12px; 
        padding: 1.5rem; 
        margin-bottom: 1.5rem; 
        border-left-width: 6px; 
        border-left-style: solid; 
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
        background: white;
    }
    .report-finding:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.12);
    }
    .risk-high { 
        background: linear-gradient(135deg, #fef2f2 0%, #fee2e2 100%); 
        border-left-color: #dc2626; 
        color: #991b1b; 
    }
    .warning-medium { 
        background: linear-gradient(135deg, #fffbeb 0%, #fef3c7 100%); 
        border-left-color: #f59e0b; 
        color: #92400e; 
    }
    .pass-ok { 
        background: linear-gradient(135deg, #f0fdf4 0%, #d1fae5 100%); 
        border-left-color: #10b981; 
        color: #065f46; 
    }
    .report-finding h4 { 
        margin: 0 0 1rem 0; 
        color: inherit; 
        font-size: 1.125rem;
        font-weight: 700;
    }
    .report-finding p { 
        margin: 0; 
        text-align: justify;
        line-height: 1.6;
        font-size: 0.95rem;
    }
    .report-finding .finding-content {
        text-align: justify;
        line-height: 1.6;
        font-size: 0.95rem;
        color: inherit;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Parse and display findings with enhanced formatting
    findings = re.split(r'(\[Risk - High\]|\[Warning - Medium\]|\[Pass\])', report_markdown)
    
    # Summary metrics
    risk_count = len([f for f in findings if f == "[Risk - High]"])
    warning_count = len([f for f in findings if f == "[Warning - Medium]"])
    pass_count = len([f for f in findings if f == "[Pass]"])
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("High Risk Issues", risk_count, delta=None)
    with col2:
        st.metric("Medium Warnings", warning_count, delta=None)
    with col3:
        st.metric("Compliance Passes", pass_count, delta=None)
    with col4:
        st.metric("Total Findings", risk_count + warning_count + pass_count, delta=None)
    
    st.divider()
    
    # Display individual findings
    for i in range(1, len(findings), 2):
        tag, content = findings[i], findings[i+1].strip()
        if not content:
            continue
            
        # Split content into title and body
        lines = content.split('\n')
        title = lines[0].strip()
        body = '\n'.join(lines[1:]).strip() if len(lines) > 1 else ""
        
        css_class, title_prefix, icon = "", "", ""
        if tag == "[Risk - High]": 
            css_class, title_prefix, icon = "risk-high", "High Risk", "⚠️"
        elif tag == "[Warning - Medium]": 
            css_class, title_prefix, icon = "warning-medium", "Medium Warning", "⚡"
        elif tag == "[Pass]": 
            css_class, title_prefix, icon = "pass-ok", "Compliance Pass", "✅"
        
        # Create enhanced finding display
        finding_html = f'''
        <div class="report-finding {css_class}">
            <h4>{icon} {title_prefix}: {title}</h4>
            <div class="finding-content">{body}</div>
        </div>
        '''
        st.markdown(finding_html, unsafe_allow_html=True)
    
    # Enhanced download section
    st.divider()
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
        padding: 1.5rem;
        border-radius: 12px;
        margin: 2rem 0;
        border: 1px solid #e2e8f0;
    ">
        <h3 style="margin: 0 0 1rem 0; color: #1e40af; text-align: center;">Download Full Report</h3>
        <p style="margin: 0; text-align: center; color: #64748b; font-size: 0.95rem;">
            Export your compliance analysis in multiple formats for sharing and documentation
        </p>
    </div>
    """, unsafe_allow_html=True)

    try:
        with st.spinner("Preparing downloads..."):
            pdf_bytes = _generate_report_pdf(report_markdown, base_filename)
            docx_bytes = _generate_report_docx(report_markdown, base_filename)
            json_bytes = _generate_report_json(report_markdown)

        # Enhanced download buttons with better styling
        col1, col2, col3, col4, col5 = st.columns(5)
        
        with col1:
            st.download_button(
                "📄 PDF Report", 
                pdf_bytes, 
                f"{base_filename}.pdf", 
                "application/pdf",
                use_container_width=True,
                help="Download as PDF for formal documentation"
            )
        with col2:
            st.download_button(
                "📝 Word Document", 
                docx_bytes, 
                f"{base_filename}.docx", 
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Download as Word document for editing"
            )
        with col3:
            st.download_button(
                "📊 JSON Data", 
                json_bytes, 
                f"{base_filename}.json", 
                "application/json",
                use_container_width=True,
                help="Download as JSON for data processing"
            )
        with col4:
            st.download_button(
                "📄 Text File", 
                report_markdown.encode('utf-8'), 
                f"{base_filename}.txt", 
                "text/plain",
                use_container_width=True,
                help="Download as plain text"
            )
        with col5:
            st.download_button(
                "📝 Markdown", 
                report_markdown.encode('utf-8'), 
                f"{base_filename}.md", 
                "text/markdown",
                use_container_width=True,
                help="Download as Markdown for version control"
            )
            
    except Exception as e:
        st.error(f"Error preparing downloads: {str(e)}")
        st.info("Please try again or contact support if the issue persists.")